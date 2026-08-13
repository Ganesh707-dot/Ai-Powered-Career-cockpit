"""Extract plain text from uploaded resume files."""

from __future__ import annotations

import io
import re
from pathlib import Path


class FileExtractError(ValueError):
    pass


ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


def detect_extension(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext in ALLOWED_EXTENSIONS:
        return ext
    if content[:5].startswith(b"%PDF"):
        return ".pdf"
    # DOCX / Office Open XML is a ZIP archive
    if content[:2] == b"PK":
        return ".docx"
    sample = content[:8000]
    if sample and b"\x00" not in sample:
        try:
            sample.decode("utf-8")
            return ".txt"
        except UnicodeDecodeError:
            pass
    raise FileExtractError(
        f"Unsupported file type '{ext or '(none)'}'. Use PDF, DOCX, TXT, or MD."
    )


def extract_text(filename: str, content: bytes) -> str:
    if not content:
        raise FileExtractError("Empty file uploaded.")

    ext = detect_extension(filename, content)

    if ext in {".txt", ".md"}:
        return content.decode("utf-8", errors="ignore").strip()

    if ext == ".pdf":
        return _extract_pdf(content)

    if ext == ".docx":
        return _extract_docx(content)

    raise FileExtractError("Unsupported file type.")


def _extract_pdf(content: bytes) -> str:
    errors: list[str] = []

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise FileExtractError(
                    "This PDF is password-protected. Export an unprotected copy or paste text."
                ) from None

        chunks: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                chunks.append(text.strip())
        joined = "\n\n".join(chunks).strip()
        if joined:
            return joined
        errors.append("pypdf returned no text")
    except FileExtractError:
        raise
    except Exception as exc:
        errors.append(f"pypdf: {exc}")

    try:
        from pdfminer.high_level import extract_text as pdfminer_extract

        mined = pdfminer_extract(io.BytesIO(content)).strip()
        mined = re.sub(r"\n{3,}", "\n\n", mined)
        if mined:
            return mined
        errors.append("pdfminer returned no text")
    except Exception as exc:
        errors.append(f"pdfminer: {exc}")

    detail = "; ".join(errors) if errors else "unknown PDF parse error"
    raise FileExtractError(
        "Could not extract text from this PDF (often a scanned/image resume). "
        "Use Paste text / statement, or export a text-based PDF. "
        f"({detail})"
    )


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise FileExtractError("DOCX support unavailable (python-docx missing).") from exc

    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    joined = "\n".join(parts).strip()
    if not joined:
        raise FileExtractError("Could not extract text from DOCX.")
    return joined
